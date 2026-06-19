import os
import sys
import time
import subprocess
import shutil
import pytest
import requests
import docx
from pathlib import Path

# Test credentials from environment for real-environment runs
TEST_EMAIL = os.getenv("TEST_EMAIL", "test@example.com")
TEST_PASSWORD = os.getenv("TEST_PASSWORD", "password123")
REPO_ROOT = Path(__file__).resolve().parents[1]
FIXTURES_DIR = REPO_ROOT / "tests" / "fixtures" / "generated"
JD_DIR = REPO_ROOT / "tests" / "jd"
BACKEND_PORT = int(os.getenv("E2E_BACKEND_PORT", "8000"))
FRONTEND_PORT = int(os.getenv("E2E_FRONTEND_PORT", "8501"))
BACKEND_URL = f"http://127.0.0.1:{BACKEND_PORT}"
FRONTEND_URL = f"http://127.0.0.1:{FRONTEND_PORT}"


def _ensure_generated_assets():
    required = [
        FIXTURES_DIR / "good_resume.docx",
        FIXTURES_DIR / "good_resume.pdf",
        FIXTURES_DIR / "bad_resume.pdf",
        FIXTURES_DIR / "resume_two_columns.pdf",
        FIXTURES_DIR / "resume_very_long_10_pages.pdf",
        JD_DIR / "small_jd.txt",
        JD_DIR / "large_jd.txt",
        JD_DIR / "unicode_jd.txt",
        JD_DIR / "ml_engineer_jd.txt",
        JD_DIR / "intern_jd.txt",
    ]
    if all(path.exists() for path in required):
        return

    from tests.generate_fixtures import main as generate_fixtures

    generate_fixtures()


def _login(page, email: str = TEST_EMAIL, password: str = TEST_PASSWORD):
    if page.locator("text=Signed in as").is_visible():
        return

    page.wait_for_selector("text=Sign in", timeout=15000)
    page.locator("input[aria-label='Email']:visible").first.fill(email)
    page.locator("input[aria-label='Password']:visible").first.fill(password)
    page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
    page.wait_for_selector(f"text=Signed in as {email}", timeout=15000)


def _logout(page):
    if page.locator("button:has-text('Sign out')").is_visible():
        page.locator("button:has-text('Sign out')").click()
        page.wait_for_selector("text=Sign in", timeout=15000)


def _fixture_path(filename: str) -> str:
    return str((FIXTURES_DIR / filename).resolve())


class AppServerController(str):
    def __new__(cls, url, backend_proc, frontend_proc, env, backend_log, frontend_log):
        obj = str.__new__(cls, url)
        obj.url = url
        obj.backend_proc = backend_proc
        obj.frontend_proc = frontend_proc
        obj.env = env
        obj.backend_log = backend_log
        obj.frontend_log = frontend_log
        return obj

    def stop_backend(self):
        if self.backend_proc:
            import platform

            if platform.system() == "Windows":
                self.backend_proc.terminate()
                try:
                    self.backend_proc.wait(timeout=5)
                except Exception:
                    subprocess.run(
                        ["taskkill", "/F", "/T", "/PID", str(self.backend_proc.pid)],
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                    )
            else:
                self.backend_proc.terminate()
                try:
                    self.backend_proc.wait(timeout=5)
                except Exception:
                    self.backend_proc.kill()
            self.backend_proc = None
            try:
                self.backend_log.close()
            except Exception:
                pass
            time.sleep(2)  # Allow port release

    def start_backend(self):
        if not self.backend_proc:
            # Re-open or append to backend log
            log_f = open(REPO_ROOT / "logs" / "backend_server.log", "a", encoding="utf-8")
            self.backend_proc = subprocess.Popen(
                [
                    sys.executable,
                    "-m",
                    "uvicorn",
                    "backend.main:app",
                    "--port",
                    str(BACKEND_PORT),
                ],
                cwd=REPO_ROOT,
                env=self.env,
                stdout=log_f,
                stderr=log_f,
            )
            self.backend_log = log_f
            # Wait/Poll backend health endpoint
            start_time = time.time()
            while time.time() - start_time < 30:
                try:
                    r = requests.get(f"{BACKEND_URL}/api/v1/health", timeout=1)
                    if r.status_code == 200:
                        break
                except Exception:
                    pass
                time.sleep(1)


@pytest.fixture(scope="module")
def app_servers():
    _ensure_generated_assets()

    # Clean up test database if it exists to ensure test isolation
    test_db = REPO_ROOT / "ats_history_test.db"
    if test_db.exists():
        try:
            test_db.unlink()
        except Exception:
            pass

    # Setup paths
    logs_dir = REPO_ROOT / "logs"
    logs_dir.mkdir(parents=True, exist_ok=True)

    backend_log_path = logs_dir / "backend_server.log"
    frontend_log_path = logs_dir / "frontend_server.log"

    # Open files for writing stdout/stderr
    backend_log = open(backend_log_path, "w", encoding="utf-8")
    frontend_log = open(frontend_log_path, "w", encoding="utf-8")
    # Build environment with MOCK_AUTH configuration
    env = os.environ.copy()
    if os.getenv("PRE_PROD", "").lower() == "true":
        env["MOCK_AUTH"] = "false"
        env.pop("ATS_FAST_MODEL_MODE", None)
    else:
        env["MOCK_AUTH"] = "true"
        env["ATS_FAST_MODEL_MODE"] = "true"
        env["GROQ_API_KEY"] = ""
    env["PYTHONUNBUFFERED"] = "1"
    env["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"
    env["STREAMLIT_SERVER_HEADLESS"] = "true"
    env["E2E_TESTING"] = "true"
    # Launch backend (uvicorn)
    backend_proc = subprocess.Popen(
        [sys.executable, "-m", "uvicorn", "backend.main:app", "--port", str(BACKEND_PORT)],
        cwd=REPO_ROOT,
        env=env,
        stdout=backend_log,
        stderr=backend_log,
    )

    # Launch frontend (streamlit)
    frontend_proc = subprocess.Popen(
        [
            sys.executable,
            "-m",
            "streamlit",
            "run",
            "frontend/streamlit_app.py",
            "--server.port",
            str(FRONTEND_PORT),
            "--server.address",
            "127.0.0.1",
        ],
        cwd=REPO_ROOT,
        env=env,
        stdout=frontend_log,
        stderr=frontend_log,
    )

    # Wait/Poll health endpoints
    start_time = time.time()
    backend_ready = False
    frontend_ready = False

    while (
        time.time() - start_time < 60
    ):  # Allow up to 60 seconds for models to load and server to start
        # Check backend
        if not backend_ready:
            try:
                r = requests.get(f"{BACKEND_URL}/api/v1/health", timeout=1)
                if r.status_code == 200:
                    backend_ready = True
            except Exception:
                pass
        # Check frontend
        if not frontend_ready:
            try:
                r = requests.get(FRONTEND_URL, timeout=1)
                if r.status_code == 200:
                    frontend_ready = True
            except Exception:
                pass

        if backend_ready and frontend_ready:
            break

        time.sleep(1)
    else:
        # Teardown processes before raising
        backend_proc.terminate()
        frontend_proc.terminate()
        backend_proc.wait()
        frontend_proc.wait()
        backend_log.close()
        frontend_log.close()
        raise RuntimeError(
            f"Servers failed to start within 60 seconds. "
            f"Backend ready: {backend_ready}, Frontend ready: {frontend_ready}. "
            f"Check logs at {backend_log_path} and {frontend_log_path}"
        )

    controller = AppServerController(
        FRONTEND_URL, backend_proc, frontend_proc, env, backend_log, frontend_log
    )
    yield controller

    # Teardown
    controller.stop_backend()
    if controller.frontend_proc:
        import platform

        if platform.system() == "Windows":
            subprocess.run(
                ["taskkill", "/F", "/T", "/PID", str(controller.frontend_proc.pid)],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        else:
            controller.frontend_proc.terminate()
        try:
            controller.frontend_proc.wait(timeout=5)
        except Exception:
            controller.frontend_proc.kill()

    # Close file handles if open
    try:
        controller.backend_log.close()
    except Exception:
        pass
    try:
        controller.frontend_log.close()
    except Exception:
        pass


def test_e2e_happy_path(app_servers, page, tmp_path):
    # 1. Open the page in browser
    page.goto(app_servers.url)

    # 2. Wait for sign-in tab/form to load
    page.wait_for_selector("text=Sign in", timeout=15000)

    # Fill in the sign-in form
    page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
    page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)

    # Submit the form
    page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()

    # 3. Verify sidebar changes to signed-in state
    page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    # 4. Click "🎯 ATS Scorer"
    page.locator("button:has-text('ATS Scorer')").click()

    # Wait for the Scorer page view
    page.locator("text=Upload your resume to begin").wait_for(state="visible", timeout=15000)

    # 5. Dynamically write a valid DOCX resume to disk
    resume_path = tmp_path / "resume.docx"
    doc = docx.Document()
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("Email: john.doe@example.com")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(
        "Senior Python Developer at Tech Corp. Developed and scaled backend APIs using FastAPI and Python."
    )
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, Docker, SQL, Git")
    doc.save(str(resume_path))

    # Upload it
    page.locator("input[type='file']").first.set_input_files(str(resume_path))

    # Wait for upload completion message in Streamlit
    page.locator("text=resume.docx").wait_for(state="visible", timeout=15000)

    # 6. Run resume analysis
    page.locator("button:has-text('Analyze Resume')").click()

    # 7. Verify that the results dashboard shows the score
    page.locator("text=Analysis Results").wait_for(state="visible", timeout=30000)
    page.locator("text=Overall ATS Score").wait_for(state="visible", timeout=30000)

    # 8. Check that the history page displays the entry
    page.locator("button:has-text('History')").click()
    page.locator("text=Analysis History").wait_for(state="visible", timeout=15000)
    page.locator("text=resume.docx").first.wait_for(state="visible", timeout=15000)


def test_e2e_error_path(app_servers, page, tmp_path):
    page.goto(app_servers.url)

    # Check if we are already signed in.
    if not page.locator("text=Signed in as").is_visible():
        page.wait_for_selector("text=Sign in", timeout=15000)
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    # Click "🎯 ATS Scorer"
    page.locator("button:has-text('ATS Scorer')").click()

    # Wait for the Scorer page view
    page.locator("text=Upload your resume to begin").wait_for(state="visible", timeout=15000)

    # Upload empty file corrupt.pdf
    corrupt_path = tmp_path / "corrupt.pdf"
    corrupt_path.write_bytes(b"")

    page.locator("input[type='file']").first.set_input_files(str(corrupt_path))
    page.locator("text=corrupt.pdf").wait_for(state="visible", timeout=15000)

    # Click analyze
    page.locator("button:has-text('Analyze Resume')").click()

    # Verify that the UI reports "Could not read or parse the resume"
    page.locator("text=Could not read or parse the resume").wait_for(state="visible", timeout=30000)


def test_login_logout(app_servers, page):
    page.goto(app_servers.url)

    _logout(page)

    # Invalid credentials should stay signed out and show an auth error.
    page.wait_for_selector("text=Sign in", timeout=15000)
    page.locator("input[aria-label='Email']:visible").first.fill("invalid@example.com")
    page.locator("input[aria-label='Password']:visible").first.fill("wrong-password")
    page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
    page.wait_for_selector("text=Wrong email or password", timeout=15000)

    # Sign-up path should create a mock session in deterministic CI mode.
    page.get_by_role("tab", name="Sign up").click()
    page.locator("input[aria-label='Email']:visible").first.fill("signup@example.com")
    page.locator("input[aria-label='Password (min 6 chars)']:visible").first.fill(TEST_PASSWORD)
    page.locator("button:has-text('Create account')").click()
    page.wait_for_selector("text=Signed in as signup@example.com", timeout=15000)

    _logout(page)
    _login(page)
    _logout(page)


def test_history(app_servers, page, tmp_path):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    resume_path = tmp_path / "history_delete_resume.docx"
    shutil.copyfile(_fixture_path("good_resume.docx"), resume_path)

    page.locator("input[type='file']").first.set_input_files(str(resume_path))
    page.locator("text=history_delete_resume.docx").wait_for(state="visible", timeout=15000)

    page.locator("button:has-text('Analyze Resume')").click()
    page.locator("text=Analysis Results").wait_for(state="visible", timeout=30000)

    # Check History
    page.locator("button:has-text('History')").click()
    page.locator("text=Analysis History").wait_for(state="visible", timeout=15000)
    page.locator("text=history_delete_resume.docx").first.wait_for(state="visible", timeout=15000)

    # Expand the history item. We wait for Streamlit to stabilize, click, and wait to verify
    page.wait_for_timeout(2000)

    # Try clicking the summary element to expand
    page.locator("[data-testid='stExpander'] summary").first.click()

    # If the delete button is still not visible after 3s, try clicking the title text
    try:
        page.locator("button:has-text('Delete')").first.wait_for(state="visible", timeout=3000)
    except Exception:
        page.locator("text=history_delete_resume.docx").first.click()
        page.wait_for_timeout(1000)

    # Delete entry
    page.locator("button:has-text('Delete')").first.click()
    page.locator("text=history_delete_resume.docx").first.wait_for(state="hidden", timeout=15000)


def test_pdf_download(app_servers, page):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    resume_path = _fixture_path("good_resume.docx")
    page.locator("input[type='file']").first.set_input_files(resume_path)
    page.locator("text=good_resume.docx").wait_for(state="visible", timeout=15000)

    page.locator("button:has-text('Analyze Resume')").click()
    page.locator("text=Analysis Results").wait_for(state="visible", timeout=30000)

    # Generate PDF
    page.locator("button:has-text('Generate PDF Report')").click()

    # Wait for download button to appear in UI
    download_btn = page.locator("[data-testid='stDownloadButton'] button:has-text('Download PDF')")
    download_btn.wait_for(state="visible", timeout=30000)

    # Expect download
    with page.expect_download() as download_info:
        download_btn.click()

    download = download_info.value
    download_path = download.path()
    assert download_path is not None
    assert os.path.exists(download_path)
    assert os.path.getsize(download_path) > 0
    assert download.suggested_filename.endswith(".pdf")


def test_invalid_resume(app_servers, page):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    resume_path = _fixture_path("bad_resume.pdf")
    page.locator("input[type='file']").first.set_input_files(resume_path)
    page.locator("text=bad_resume.pdf").wait_for(state="visible", timeout=15000)

    page.locator("button:has-text('Analyze Resume')").click()
    page.locator("text=Could not read or parse the resume").wait_for(state="visible", timeout=15000)


def test_large_resume(app_servers, page, tmp_path):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    # Generate >5MB file to trigger limit
    huge_file = tmp_path / "huge_resume.pdf"
    huge_file.write_bytes(b"%PDF-1.4\n" + b"0" * (6 * 1024 * 1024))

    page.locator("input[type='file']").first.set_input_files(str(huge_file))
    page.locator("text=huge_resume.pdf").wait_for(state="visible", timeout=15000)

    # Click analyze to trigger backend size check
    page.locator("button:has-text('Analyze Resume')").click()

    # Verify limit error in UI
    page.locator("text=File too large").wait_for(state="visible", timeout=15000)


def test_multiple_uploads(app_servers, page):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    # First run
    resume_path_1 = _fixture_path("good_resume.docx")
    page.locator("input[type='file']").first.set_input_files(resume_path_1)
    page.locator("text=good_resume.docx").wait_for(state="visible", timeout=15000)
    page.locator("button:has-text('Analyze Resume')").click()
    page.locator("text=Analysis Results").wait_for(state="visible", timeout=30000)

    # Second run
    resume_path_2 = _fixture_path("good_resume.pdf")
    page.locator("input[type='file']").first.set_input_files(resume_path_2)
    page.locator("text=good_resume.pdf").wait_for(state="visible", timeout=15000)
    page.locator("button:has-text('Analyze Resume')").click()
    page.locator("text=Analysis Results").wait_for(state="visible", timeout=30000)


def test_network_failure(app_servers, page):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    resume_path = _fixture_path("good_resume.docx")
    page.locator("input[type='file']").first.set_input_files(resume_path)
    page.locator("text=good_resume.docx").wait_for(state="visible", timeout=15000)

    # Stop backend to simulate network failure
    app_servers.stop_backend()

    page.locator("button:has-text('Analyze Resume')").click()

    # Verify friendly error is shown and UI recovers
    page.locator("text=Could not reach the backend").wait_for(state="visible", timeout=15000)

    # Restart backend
    app_servers.start_backend()


def test_server_restart(app_servers, page):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    resume_path = _fixture_path("good_resume.docx")
    page.locator("input[type='file']").first.set_input_files(resume_path)
    page.locator("text=good_resume.docx").wait_for(state="visible", timeout=15000)

    # Stop backend
    app_servers.stop_backend()

    # Attempt analyze -> should show reachability error
    page.locator("button:has-text('Analyze Resume')").click()
    page.locator("text=Could not reach").wait_for(state="visible", timeout=15000)

    # Restart backend
    app_servers.start_backend()

    # Retry -> should succeed
    page.locator("button:has-text('Analyze Resume')").click()
    page.locator("text=Analysis Results").wait_for(state="visible", timeout=30000)


def test_e2e_matrix(app_servers, page):
    page.goto(app_servers.url)

    if not page.locator("text=Signed in as").is_visible():
        page.locator("input[aria-label='Email']:visible").first.fill(TEST_EMAIL)
        page.locator("input[aria-label='Password']:visible").first.fill(TEST_PASSWORD)
        page.locator("[data-testid='stFormSubmitButton'] button:visible").first.click()
        page.wait_for_selector(f"text=Signed in as {TEST_EMAIL}", timeout=15000)

    page.locator("button:has-text('ATS Scorer')").click()

    # 5 matrix test combinations: (resume_filename, jd_filename)
    import shutil

    has_ocr = shutil.which("tesseract") is not None and shutil.which("pdftoppm") is not None

    matrix = [
        ("good_resume.docx", "small_jd.txt"),
        ("unicode_resume.docx", "unicode_jd.txt"),
        ("resume_scanned_image.pdf" if has_ocr else "good_resume.pdf", "ml_engineer_jd.txt"),
        ("resume_two_columns.pdf", "large_jd.txt"),
        ("resume_very_long_10_pages.pdf", "intern_jd.txt"),
    ]

    for resume_name, jd_name in matrix:
        # Switch/Enable Job Description Comparison Mode using precise radio locator
        page.locator("[data-testid='stRadio']").get_by_text(
            "Job Description Comparison", exact=True
        ).click()

        # Input the JD text
        jd_path = JD_DIR / jd_name
        jd_content = jd_path.read_text(encoding="utf-8")

        page.locator("textarea[placeholder='Paste the JD here...']").first.fill(jd_content)

        # Upload the resume file
        resume_path = _fixture_path(resume_name)
        page.locator("input[type='file']").first.set_input_files(resume_path)
        page.locator(f"text={resume_name}").wait_for(state="visible", timeout=15000)

        # Run analysis
        page.locator("button:has-text('Analyze Resume')").click()

        # Verify results dashboard
        page.locator("text=Analysis Results").wait_for(state="visible", timeout=30000)


def test_e2e_artifact_generation_on_failure(tmp_path):
    # Dummy failing test to verify that the playwright_artifacts fixture collects screenshot and trace
    dummy_test_file = tmp_path / "test_dummy_failure_for_artifacts.py"
    dummy_test_file.write_text("""
def test_dummy_fail_run(page):
    page.goto("about:blank")
    assert False
""")

    import shutil

    shutil.copy(REPO_ROOT / "tests" / "conftest.py", tmp_path / "conftest.py")

    artifacts_dir = REPO_ROOT / "artifacts"
    if artifacts_dir.exists():
        for f in artifacts_dir.glob("failure_test_dummy_fail_run*.png"):
            try:
                f.unlink()
            except Exception:
                pass
        for f in artifacts_dir.glob("trace_test_dummy_fail_run*.zip"):
            try:
                f.unlink()
            except Exception:
                pass

    res = subprocess.run(
        [sys.executable, "-m", "pytest", str(dummy_test_file), "-v"],
        capture_output=True,
        text=True,
        cwd=REPO_ROOT,
    )

    assert res.returncode != 0

    screenshots = list(artifacts_dir.glob("failure_test_dummy_fail_run*.png"))
    traces = list(artifacts_dir.glob("trace_test_dummy_fail_run*.zip"))

    assert (
        len(screenshots) > 0
    ), f"Failure screenshot not created! Output: {res.stdout}\n{res.stderr}"
    assert len(traces) > 0, f"Failure trace not created! Output: {res.stdout}\n{res.stderr}"

    for f in screenshots:
        try:
            f.unlink()
        except Exception:
            pass
    for f in traces:
        try:
            f.unlink()
        except Exception:
            pass
