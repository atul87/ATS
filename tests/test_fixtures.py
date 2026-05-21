import io
from pathlib import Path

import pytest

from backend.services.resume_parser import parse_resume_file, FileParsingError


def _ensure_fixture_dir():
    d = Path("tests/fixtures")
    d.mkdir(parents=True, exist_ok=True)
    return d


def _make_docx(text: str, filename: str) -> bytes:
    from docx import Document

    doc = Document()
    if text:
        for line in text.split("\n"):
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    data = bio.getvalue()
    # also write to disk for inspection
    p = _ensure_fixture_dir() / filename
    p.write_bytes(data)
    return data


def _make_pdf(text: str, filename: str, image_only: bool = False) -> bytes:
    # Create a simple PDF using reportlab. If image_only, embed an image instead of text.
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    from PIL import Image, ImageDraw

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    width, height = letter

    if image_only:
        # create an image and draw some text into it (so it is not extractable text)
        img = Image.new("RGB", (int(width), int(height)), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.text((50, 100), "This is an image-only resume (Python, Docker)", fill=(0, 0, 0))
        img_bio = io.BytesIO()
        img.save(img_bio, format="PNG")
        img_bio.seek(0)
        c.drawImage(ImageReader(img_bio), 0, 0, width=width, height=height)
    else:
        # Draw selectable text
        text_lines = text.split("\n")
        y = height - 100
        for line in text_lines:
            c.drawString(50, y, line)
            y -= 14

    c.showPage()
    c.save()
    data = bio.getvalue()
    p = _ensure_fixture_dir() / filename
    p.write_bytes(data)
    return data


class ImageReader:
    # tiny adapter so reportlab can accept a BytesIO image
    def __init__(self, bio):
        self._bio = bio

    def getSize(self):
        from PIL import Image

        self._bio.seek(0)
        im = Image.open(self._bio)
        return im.size

    def _image(self):
        self._bio.seek(0)
        return self._bio


def test_fixtures_parse_and_failures():
    fixtures = _ensure_fixture_dir()

    # normal docx
    normal_docx = _make_docx("John Doe\nSkills: Python, Docker, SQL", "normal_resume.docx")
    text, meta = parse_resume_file(normal_docx, "normal_resume.docx")
    assert meta["success"] is True
    assert "python" in text.lower()

    # unicode docx
    uni_docx = _make_docx("名字: 张伟\n技能: Python, 数据库", "unicode_resume.docx")
    text_u, meta_u = parse_resume_file(uni_docx, "unicode_resume.docx")
    assert meta_u["success"] is True
    assert len(text_u) > 0

    # table-like docx
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Company"
    table.cell(1, 0).text = "Software Engineer"
    table.cell(1, 1).text = "Example Inc"
    bio = io.BytesIO()
    doc.save(bio)
    table_bytes = bio.getvalue()
    (fixtures / "table_resume.docx").write_bytes(table_bytes)
    text_t, meta_t = parse_resume_file(table_bytes, "table_resume.docx")
    assert meta_t["success"] is True
    assert "software engineer" in text_t.lower()

    # normal selectable PDF
    pdf_bytes = _make_pdf("Resume\nSkills: Python, FastAPI", "normal_resume.pdf")
    text_p, meta_p = parse_resume_file(pdf_bytes, "normal_resume.pdf")
    assert meta_p["success"] is True
    assert "python" in text_p.lower()

    # image-only PDF should fail text extraction
    img_pdf = _make_pdf("", "image_only_resume.pdf", image_only=True)
    with pytest.raises(FileParsingError):
        parse_resume_file(img_pdf, "image_only_resume.pdf")

    # malformed PDF
    malformed = b"%PDF-1.4\n%corrupt\n" + b"0" * 10
    (fixtures / "malformed.pdf").write_bytes(malformed)
    with pytest.raises(FileParsingError):
        parse_resume_file(malformed, "malformed.pdf")

    # empty docx
    empty = _make_docx("", "empty.docx")
    with pytest.raises(FileParsingError):
        parse_resume_file(empty, "empty.docx")
