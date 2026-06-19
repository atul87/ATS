import io
from pathlib import Path
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image, ImageDraw


def ensure_dirs():
    fixtures_dir = Path("tests/fixtures/generated")
    jd_dir = Path("tests/jd")
    fixtures_dir.mkdir(parents=True, exist_ok=True)
    jd_dir.mkdir(parents=True, exist_ok=True)
    return fixtures_dir, jd_dir


def create_good_resume_docx(path):
    doc = Document()
    doc.add_heading("John Doe - Senior Software Engineer", 0)
    doc.add_paragraph(
        "Email: john.doe@example.com | Phone: 123-456-7890 | GitHub: github.com/johndoe"
    )
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Highly motivated Senior Software Engineer with 8 years of experience building scalable web applications. Expert in Python, FastAPI, Docker, and PostgreSQL."
    )
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Lead Developer at Tech Corp (Jan 2020 - Present)")
    doc.add_paragraph(
        "Designed and implemented microservices using FastAPI and Docker. Scaled API performance by 40%. Managed a team of 4 engineers."
    )
    doc.add_paragraph("Software Engineer at Web Inc (Jun 2017 - Dec 2019)")
    doc.add_paragraph(
        "Developed backend APIs using Django and Python. Wrote SQL queries and optimized database schemas."
    )
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Science in Computer Science - State University, 2017")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, Django, Docker, Kubernetes, PostgreSQL, SQL, Git, AWS")
    doc.save(path)


def create_good_resume_pdf(path):
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 50
    lines = [
        "Jane Smith - Senior Cloud Architect",
        "Email: jane.smith@example.com | LinkedIn: linkedin.com/in/janesmith",
        "",
        "Professional Summary:",
        "Experienced Cloud Architect specializing in cloud infrastructure, containerization, and backend APIs.",
        "",
        "Experience:",
        "Principal Architect - CloudScale Systems (Mar 2021 - Present)",
        "- Led migration of legacy applications to AWS using Kubernetes and Docker.",
        "- Implemented high-throughput REST APIs using Python and FastAPI.",
        "Senior Developer - DevLabs (Sep 2018 - Feb 2021)",
        "- Built backend microservices with Python, PostgreSQL, and Git.",
        "",
        "Skills:",
        "Python, FastAPI, AWS, Kubernetes, Docker, PostgreSQL, Git, Linux",
    ]
    for line in lines:
        c.drawString(50, y, line)
        y -= 15
    c.showPage()
    c.save()


def create_empty_pdf(path):
    # Generates a valid PDF structure but with no text content
    c = canvas.Canvas(str(path), pagesize=letter)
    c.showPage()
    c.save()


def create_image_resume_pdf(path):
    # PDF containing an image of text (requires OCR)
    width, height = letter
    img = Image.new("RGB", (int(width), int(height)), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((50, 100), "OCR Scanned Resume", fill=(0, 0, 0))
    draw.text((50, 120), "Name: OCR Candidate", fill=(0, 0, 0))
    draw.text((50, 140), "Skills: Python, Docker, PyTorch, SQL", fill=(0, 0, 0))
    draw.text((50, 160), "Experience in Machine Learning and Deep Learning.", fill=(0, 0, 0))

    img_bio = io.BytesIO()
    img.save(img_bio, format="PNG")
    img_bio.seek(0)

    c = canvas.Canvas(str(path), pagesize=letter)
    from reportlab.lib.utils import ImageReader

    c.drawImage(ImageReader(img_bio), 0, 0, width=width, height=height)
    c.showPage()
    c.save()


def create_unicode_resume_docx(path):
    doc = Document()
    doc.add_heading("张伟 - 软件工程师", 0)
    doc.add_paragraph("邮箱: zhangwei@example.com | 电话: +86 10 1234 5678")
    doc.add_heading("专业技能", level=1)
    doc.add_paragraph("编程语言: Python, Go, C++")
    doc.add_paragraph("技术框架: FastAPI, Django, Flask, TensorFlow")
    doc.add_paragraph("数据库: PostgreSQL, Redis, MySQL")
    doc.add_heading("工作经历", level=1)
    doc.add_paragraph("高级开发工程师 - 北京科技有限公司 (2020年至今)")
    doc.add_paragraph(
        "使用Python和FastAPI开发高并发后端服务。负责容器化部署 (Docker, Kubernetes)。"
    )
    doc.save(path)


def create_large_resume_pdf(path):
    # Generate a large PDF file close to the limit of 5MB (e.g. ~4.6 MB)
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter

    # Write some selectable text
    c.drawString(50, height - 50, "Large Resume Document")
    c.drawString(50, height - 70, "Skills: Python, FastAPI, Docker, SQL, Kubernetes")

    # Embed a very large, high-res dummy image to bloat the file size
    large_img = Image.new("RGB", (3000, 3000), color=(240, 240, 240))
    img_bio = io.BytesIO()
    # Save with no compression to ensure large size
    large_img.save(img_bio, format="PNG")
    img_bio.seek(0)

    from reportlab.lib.utils import ImageReader

    c.drawImage(ImageReader(img_bio), 50, 50, width=width - 100, height=height - 200)
    c.showPage()
    c.save()


def create_bad_resume_pdf(path):
    # Corrupted/invalid PDF file
    path.write_bytes(b"%PDF-1.4\n%THIS IS A CORRUPT RESUME FILE BYTES\n" + b"A" * 1000)


def create_table_resume_docx(path):
    doc = Document()
    doc.add_heading("Table-structured Resume", 0)
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Personal Details"
    table.cell(0, 1).text = "Alex Johnson | alex@example.com"
    table.cell(1, 0).text = "Experience"
    table.cell(1, 1).text = "Senior Python Developer at Tech Group (FastAPI, Docker, SQL)"
    table.cell(2, 0).text = "Education"
    table.cell(2, 1).text = "MS in Computer Science, 2019"
    table.cell(3, 0).text = "Skills"
    table.cell(3, 1).text = "Python, FastAPI, Docker, AWS, PostgreSQL, Git"
    doc.save(path)


def create_multilingual_resume_pdf(path):
    c = canvas.Canvas(str(path), pagesize=letter)
    _, height = letter
    y = height - 50
    lines = [
        "Multilingual Resume - Carlos Müller",
        "Email: carlos@example.com",
        "",
        "English Summary:",
        "Bilingual Backend Developer experienced in Python, FastAPI and SQL.",
        "",
        "Deutsch (German) Summary:",
        "Zweisprachiger Backend-Entwickler mit Erfahrung in Python, FastAPI und SQL.",
        "",
        "Español (Spanish) Summary:",
        "Desarrollador backend bilingüe con experiencia en Python, FastAPI y SQL.",
        "",
        "Skills:",
        "Python, FastAPI, SQL, Docker, Spanish, German, English",
    ]
    for line in lines:
        c.drawString(50, y, line)
        y -= 15
    c.showPage()
    c.save()


def create_resume_with_tables_pdf(path):
    # PDF resume with custom table grids using reportlab canvas
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    c.drawString(50, height - 50, "Resume with Table Grid Layout")

    # Draw simple table border lines
    c.rect(50, height - 200, width - 100, 120, fill=0, stroke=1)
    c.line(150, height - 200, 150, height - 80)
    c.line(50, height - 140, width - 50, height - 140)

    c.drawString(60, height - 110, "Name")
    c.drawString(160, height - 110, "Sarah Connor (Python, Docker)")
    c.drawString(60, height - 170, "Skills")
    c.drawString(160, height - 170, "Python, Django, SQL, PyTorch, Kubernetes, Git")
    c.showPage()
    c.save()


def create_resume_two_columns_pdf(path):
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter

    # Header
    c.drawString(50, height - 50, "Robert Downey - Backend Engineer")
    c.drawString(50, height - 65, "robert@example.com | 415-555-0199")
    c.line(50, height - 75, width - 50, height - 75)

    # Left Column (e.g. Contact, Skills, Languages)
    c.drawString(50, height - 100, "SKILLS")
    c.drawString(50, height - 115, "Python")
    c.drawString(50, height - 130, "FastAPI")
    c.drawString(50, height - 145, "Docker")
    c.drawString(50, height - 160, "PostgreSQL")
    c.drawString(50, height - 175, "AWS")

    # Right Column (e.g. Experience, Projects)
    c.drawString(250, height - 100, "PROFESSIONAL EXPERIENCE")
    c.drawString(250, height - 115, "Lead Software Engineer - Marvel Tech (2018 - Present)")
    c.drawString(250, height - 130, "Built robust web backend services using Python and FastAPI.")
    c.drawString(250, height - 145, "Architected Docker-based container deployments on AWS.")

    c.showPage()
    c.save()


def create_resume_scanned_image_pdf(path):
    # Another scanned image PDF with different text and format
    width, height = letter
    img = Image.new("RGB", (int(width), int(height)), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((50, 100), "Scanned Profile - Image Format Only", fill=(0, 0, 0))
    draw.text((50, 120), "Role: Machine Learning Researcher", fill=(0, 0, 0))
    draw.text((50, 140), "Primary Tooling: Python, PyTorch, Scikit-learn, SQL", fill=(0, 0, 0))

    img_bio = io.BytesIO()
    img.save(img_bio, format="PNG")
    img_bio.seek(0)

    c = canvas.Canvas(str(path), pagesize=letter)
    from reportlab.lib.utils import ImageReader

    c.drawImage(ImageReader(img_bio), 0, 0, width=width, height=height)
    c.showPage()
    c.save()


def create_resume_emoji_docx(path):
    doc = Document()
    doc.add_heading("🚀 Developer Extraordinaire 🚀", 0)
    doc.add_paragraph("📧 email: emoji@example.com | 📱 phone: 123-emoji")
    doc.add_heading("💼 Experience", level=1)
    doc.add_paragraph("🌟 Senior Engineer at EmojiLabs (2021 - Present)")
    doc.add_paragraph(
        "💻 Designed high-availability FastAPI microservices. ⚡ Achieved 99.9% uptime using Docker container clustering."
    )
    doc.add_heading("🛠️ Skills", level=1)
    doc.add_paragraph("Python, FastAPI, SQL, Docker, AWS, Git, 🤖 Machine Learning")
    doc.save(path)


def create_resume_very_long_pdf(path):
    c = canvas.Canvas(str(path), pagesize=letter)
    _, height = letter
    # We will write 10 pages to simulate a very long resume document
    for page_num in range(1, 11):
        c.drawString(50, height - 50, f"Dr. Albert Einstein - Resume Page {page_num} of 10")
        c.drawString(50, height - 80, "Research & Development Portfolio")
        c.drawString(
            50, height - 110, f"Detailed academic publication logs for section {page_num}:"
        )

        # Write some dummy lines
        y = height - 140
        for i in range(1, 20):
            c.drawString(
                50,
                y,
                f"Publication #{page_num}.{i}: Advanced Python, FastAPI & Docker in Large-Scale Infrastructure.",
            )
            y -= 25

        c.drawString(50, 50, "Skills: Python, FastAPI, Docker, SQL, PyTorch, Research")
        c.showPage()
    c.save()


def create_jd_fixtures(jd_dir):
    # 1. small_jd.txt
    (jd_dir / "small_jd.txt").write_text(
        "Python Developer\nRequired: Python, FastAPI, SQL.\n", encoding="utf-8"
    )

    # 2. large_jd.txt
    large_jd_content = (
        "Position: Enterprise Solutions Architect / Backend Lead Engineer\n\n"
        "We are looking for an exceptional engineer to design and scale our cloud services.\n"
        "Key Responsibilities:\n"
        "1. Architect microservices using Python and FastAPI.\n"
        "2. Dockerize applications and manage deployments on Kubernetes.\n"
        "3. Manage and optimize large relational databases (PostgreSQL/SQL).\n"
        "4. Collaborate with cross-functional teams to integrate AWS cloud tools.\n"
        "Requirements:\n"
        "- 7+ years of experience in backend development.\n"
        "- Strong proficiency in Python, Django, FastAPI, and SQL.\n"
        "- Experience with Docker, Git, CI/CD pipelines, and AWS/Kubernetes.\n"
        "- Experience with Machine Learning pipelines is a plus.\n"
        "- Excellent communication and mentoring skills.\n"
    )
    (jd_dir / "large_jd.txt").write_text(large_jd_content, encoding="utf-8")

    # 3. empty_jd.txt
    (jd_dir / "empty_jd.txt").write_text("", encoding="utf-8")

    # 4. unicode_jd.txt
    (jd_dir / "unicode_jd.txt").write_text(
        "职位: 软件开发工程师\n要求技能: Python, FastAPI, 数据库 SQL, Docker。\n", encoding="utf-8"
    )

    # 5. ml_engineer_jd.txt
    ml_jd = (
        "Machine Learning Engineer\n\n"
        "Responsibilities:\n"
        "- Build, train and deploy ML/DL models in production.\n"
        "- Perform data analysis and feature engineering.\n"
        "Requirements:\n"
        "- Strong programming in Python.\n"
        "- Hands-on experience with PyTorch, TensorFlow, Scikit-learn, and NumPy.\n"
        "- Database knowledge in SQL/PostgreSQL.\n"
        "- Experience with Docker for ML model serving.\n"
    )
    (jd_dir / "ml_engineer_jd.txt").write_text(ml_jd, encoding="utf-8")

    # 6. intern_jd.txt
    intern_jd = (
        "Software Engineering Intern\n\n"
        "Requirements:\n"
        "- Basic knowledge of programming, preferably Python.\n"
        "- Experience with Git and version control.\n"
        "- Eagerness to learn new technologies like FastAPI, Docker, and SQL.\n"
    )
    (jd_dir / "intern_jd.txt").write_text(intern_jd, encoding="utf-8")


def main():
    fixtures_dir, jd_dir = ensure_dirs()
    print("Generating resume fixtures...")
    create_good_resume_docx(fixtures_dir / "good_resume.docx")
    create_good_resume_pdf(fixtures_dir / "good_resume.pdf")
    create_empty_pdf(fixtures_dir / "empty.pdf")
    create_image_resume_pdf(fixtures_dir / "image_resume.pdf")
    create_unicode_resume_docx(fixtures_dir / "unicode_resume.docx")
    create_large_resume_pdf(fixtures_dir / "large_resume.pdf")
    create_bad_resume_pdf(fixtures_dir / "bad_resume.pdf")
    create_table_resume_docx(fixtures_dir / "table_resume.docx")
    create_multilingual_resume_pdf(fixtures_dir / "multilingual_resume.pdf")

    # Expanded fixtures
    create_resume_with_tables_pdf(fixtures_dir / "resume_with_tables.pdf")
    create_resume_two_columns_pdf(fixtures_dir / "resume_two_columns.pdf")
    create_resume_scanned_image_pdf(fixtures_dir / "resume_scanned_image.pdf")
    create_resume_emoji_docx(fixtures_dir / "resume_emoji.docx")
    create_resume_very_long_pdf(fixtures_dir / "resume_very_long_10_pages.pdf")

    print("Generating job description fixtures...")
    create_jd_fixtures(jd_dir)
    print("All fixtures generated successfully!")


if __name__ == "__main__":
    main()
