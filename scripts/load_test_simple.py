import os
import time
import io
import statistics
from concurrent.futures import ThreadPoolExecutor, as_completed
import requests
from docx import Document

BACKEND = os.environ.get("BACKEND_URL", "http://127.0.0.1:8000")
URL = BACKEND.rstrip("/") + "/api/v1/analyze-resume"
TOKEN = os.environ.get("LOAD_TEST_TOKEN", "mock-access-token")
TOTAL = int(os.environ.get("LOAD_TOTAL", "200"))
CONCURRENCY = int(os.environ.get("LOAD_CONCURRENCY", "50"))
TIMEOUT = int(os.environ.get("LOAD_TIMEOUT", "30"))


def make_docx_bytes():
    doc = Document()
    doc.add_heading("Jane Doe", 0)
    doc.add_paragraph("jane@example.com")
    doc.add_paragraph("Skills: Python, FastAPI, Docker, SQL, AWS")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def post_resume(session, data_bytes):
    files = {
        "resume": (
            "load_resume.docx",
            data_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    headers = {"Authorization": f"Bearer {TOKEN}"}
    start = time.time()
    try:
        r = session.post(
            URL,
            files=files,
            data={"job_description": "Python, FastAPI"},
            headers=headers,
            timeout=TIMEOUT,
        )
        elapsed = (time.time() - start) * 1000
        return (r.status_code, elapsed, None if r.status_code == 200 else r.text[:200])
    except Exception as e:
        elapsed = (time.time() - start) * 1000
        return (None, elapsed, str(e))


def run_load():
    data = make_docx_bytes()
    latencies = []
    failures = 0
    statuses = {}
    with ThreadPoolExecutor(max_workers=CONCURRENCY) as ex:
        futures = []
        session = requests.Session()
        for i in range(TOTAL):
            futures.append(ex.submit(post_resume, session, data))
        for f in as_completed(futures):
            status, elapsed, err = f.result()
            latencies.append(elapsed)
            statuses[status] = statuses.get(status, 0) + 1
            if status != 200:
                failures += 1
    if latencies:
        p95 = statistics.quantiles(latencies, n=100)[94]
    else:
        p95 = 0
    print(f"Requests: {TOTAL}, Concurrency: {CONCURRENCY}")
    print(f"Status counts: {statuses}")
    print(f"Failures: {failures} ({failures/TOTAL:.2%})")
    print(f"p95 latency: {p95:.0f} ms")


if __name__ == "__main__":
    run_load()
